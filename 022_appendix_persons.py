import csv
import os
from docx import Document
from docx.oxml import OxmlElement
from copy import deepcopy
from datetime import datetime
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt  # Für präzise Abstände


def erstelle_personen_anhang():
    print("--- Personen-Anhang Generator ---")
    input_csv = input("Bitte den Pfad zur CSV-Datei eingeben: ").strip().strip('"')

    if not input_csv:
        input_csv = r"D:\heiBOX\Seafile\Masterarbeit_Ablage\daten_heibox\260117_PERSONEN.csv"
        print(f"Keine Eingabe erkannt. Nutze Standardpfad: {input_csv}")

    zeitstempel = datetime.now().strftime("%y%m%d_%H%M")
    dateiname = f"{zeitstempel}_Anhang_Personen.docx"

    template_path = r"D:\heiBOX\Seafile\Masterarbeit_Ablage\anhang_heibox\template_appendix_person.docx"
    output_dir = r"D:\heiBOX\Seafile\Masterarbeit_Ablage\anhang_heibox"
    output_docx = os.path.join(output_dir, dateiname)

    if not os.path.exists(output_dir):
        os.makedirs(output_dir)

    if not os.path.exists(template_path) or not os.path.exists(input_csv):
        print("\nFehler: Pfade prüfen!")
        return

    doc = Document(template_path)

    # --- GLOBALER ZEILENABSTAND (1,15) ---
    style = doc.styles['Normal']
    style.paragraph_format.line_spacing = 1.15

    # Vorlage-Elemente extrahieren
    template_para = doc.paragraphs[0]
    template_table = doc.tables[0]

    template_para_xml = deepcopy(template_para._element)
    template_table_xml = deepcopy(template_table._element)

    spalten_liste = [
        "KORR-NACHNAME", "KORR-VORNAME", "KORR-BERU",
        "TAG-TAET", "TAG-FACH", "TAG-INST",
        "REF-1", "REF-2", "REF-3", "REF-4"
    ]

    try:
        with open(input_csv, mode='r', encoding='utf-8-sig') as csvfile:
            sample = csvfile.read(4096)
            csvfile.seek(0)
            dialect = csv.Sniffer().sniff(sample)
            reader = csv.DictReader(csvfile, dialect=dialect)

            template_para._element.getparent().remove(template_para._element)
            template_table._element.getparent().remove(template_table._element)

            i = 0
            for i, row in enumerate(reader, start=1):
                # A) Absatz einfügen
                new_p_element = deepcopy(template_para_xml)
                doc._element.body.append(new_p_element)
                aktiver_absatz = doc.paragraphs[-1]

                # Zeilenabstand auch für den neuen Absatz sicherstellen
                aktiver_absatz.paragraph_format.line_spacing = 1.15

                for spalte in spalten_liste:
                    if spalte in aktiver_absatz.text:
                        wert = str(row.get(spalte, "")).strip()
                        aktiver_absatz.text = aktiver_absatz.text.replace(spalte, wert)

                # B) Tabelle einfügen
                new_tbl_element = deepcopy(template_table_xml)
                doc._element.body.append(new_tbl_element)
                aktuelle_tabelle = doc.tables[-1]

                for zeile in aktuelle_tabelle.rows:
                    for zelle in zeile.cells:
                        cell_text = zelle.text
                        for spalte in spalten_liste:
                            if spalte in cell_text:
                                wert = str(row.get(spalte, "")).strip()
                                cell_text = cell_text.replace(spalte, wert)

                        # Leere Zeilen in Referenzen löschen
                        lines = cell_text.splitlines()
                        cleaned_lines = [l.strip() for l in lines if l.strip() and any(c.isalnum() for c in l)]
                        zelle.text = "\n".join(cleaned_lines)

                        # Linksbündig & 1,15 Zeilenabstand pro Zelle
                        for paragraph in zelle.paragraphs:
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                            paragraph.paragraph_format.line_spacing = 1.15

                # C) Leerzeile
                doc._element.body.append(OxmlElement('w:p'))

        doc.save(output_docx)
        print(f"\nErfolg! {i} Personen wurden verarbeitet.")
        print(f"Datei: {dateiname}")

    except Exception as e:
        print(f"\nFehler: {e}")


if __name__ == "__main__":
    erstelle_personen_anhang()