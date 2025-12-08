import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
import os
import sys

# --- 1. Konfiguration ---
# Fester Pfad zum Eingabeordner
INPUT_FOLDER = r'D:\heiBOX\Seafile\Masterarbeit_Ablage\daten_heibox'

# Basisname für die Ausgabedatei
OUTPUT_FILENAME_BASE = 'Alle_Briefe_Chronologisch_Final.docx'

# --- 2. SPALTENZUORDNUNG ---
COLUMN_MAPPING = {
    'BRIEF-ID': 'ID',
    'BRIEF-TITEL': 'Title',
    'DATUM': 'Date',
    'ABS-ORT': 'Versandort',
    'EMP-ORT': 'Empfangsort',
    'TRANSK': 'Text_Full',
    'BESCHR': 'Description',
    'TAG-FUNK': 'Funktionstags',
    'NONOS-TAET': 'Tätigkeit',
    'NONOS-FACH': 'Fach',
    'NONOS-INST': 'Institution'
}


# -------------------------------------------------------------------------------------

# Hilfsfunktionen bleiben unverändert
def get_formatted_date_string(date_value):
    """Konvertiert ein pandas Timestamp-Objekt sicher in YYYY-MM-DD."""
    if pd.notna(date_value) and pd.api.types.is_datetime64_any_dtype(date_value):
        try:
            return date_value.strftime('%Y-%m-%d')
        except:
            return 'Datum ungültig'
    return 'Datum unbekannt'


def add_metadata(doc, key, value):
    if pd.notna(value) and value:
        if key == "Datum" and pd.api.types.is_datetime64_any_dtype(value):
            display_value = value.strftime('%Y-%m-%d')
        else:
            display_value = str(value)
    else:
        display_value = 'N/A'

    p = doc.add_paragraph()
    p.add_run(f"{key}: ").bold = True
    p.add_run(display_value)


# 3. Interaktive Abfrage und Pfadkonstruktion
try:
    csv_filename = input("Bitte geben Sie den Namen Ihrer CSV-Datei ein (z.B. daten.csv): ")

    # Vollständigen Pfad zur CSV-Datei erstellen
    CSV_FILE_PATH = os.path.join(INPUT_FOLDER, csv_filename)

    # NEU: Pfad für die Ausgabedatei im selben Ordner erstellen
    OUTPUT_FILE_PATH = os.path.join(INPUT_FOLDER, OUTPUT_FILENAME_BASE)

    print(f"Versuche Datei zu lesen von: {CSV_FILE_PATH}")
    print(f"Ausgabedokument wird gespeichert unter: {OUTPUT_FILE_PATH}")

except Exception as e:
    print(f"Fehler bei der Pfadkonstruktion: {e}")
    sys.exit()

# 4. CSV-Datei laden, Spalten umbenennen und sortieren
try:
    df = pd.read_csv(CSV_FILE_PATH)

    missing_cols = [col for col in COLUMN_MAPPING.keys() if col not in df.columns]
    if missing_cols:
        print(f"FEHLER: Die folgenden Spalten wurden in der CSV-Datei nicht gefunden:")
        print(missing_cols)
        print("\nBitte überprüfen Sie die Spaltennamen in der COLUMN_MAPPING im Skript.")
        sys.exit()

    df.rename(columns=COLUMN_MAPPING, inplace=True)

    # DATUMSKONVERTIERUNG UND SORTIERUNG
    df['Date'] = pd.to_datetime(df['Date'], errors='coerce')
    df.sort_values(by='Date', ascending=True, inplace=True)

    num_unsorted = df['Date'].isna().sum()
    if num_unsorted > 0:
        print(f"Hinweis: {num_unsorted} Einträge mit ungültigem Datum wurden ans Ende der Liste gesetzt.")

except FileNotFoundError:
    print(f"\nFATALER FEHLER: Die Datei '{CSV_FILE_PATH}' konnte nicht gefunden werden.")
    print("Bitte stellen Sie sicher, dass der Dateiname korrekt ist und die Datei im Ordner existiert.")
    sys.exit()
except Exception as e:
    print(f"Ein unerwarteter Fehler ist aufgetreten: {e}")
    sys.exit()

# 5. Word-Dokument initialisieren
document = Document()
print(f"Es wurden {len(df)} Einträge geladen und chronologisch sortiert. Die Dokumenterstellung beginnt...")

# 6. Iteriere über jede Zeile (chronologisch) und erstelle eine Seite
for index, row in df.iterrows():
    # --- Start des Briefinhalts / der Seite ---

    # 6.1. Hauptüberschrift
    title_text = row.get('Title', 'Unbetitelter Brief')
    date_display_str = get_formatted_date_string(row.get('Date'))

    document.add_heading(f"Brief ({date_display_str}) – {title_text}", level=1)

    # 6.2. Metadaten-Block
    document.add_heading("Details", level=3)

    add_metadata(document, "ID", row.get('ID'))
    add_metadata(document, "Datum", row.get('Date'))
    add_metadata(document, "Versandort", row.get('Versandort'))
    add_metadata(document, "Empfangsort", row.get('Empfangsort'))
    add_metadata(document, "Beschreibung", row.get('Description'))
    add_metadata(document, "Funktionstags", row.get('Funktionstags'))
    document.add_paragraph("--- Personen Details ---")
    add_metadata(document, "Tätigkeit", row.get('Tätigkeit'))
    add_metadata(document, "Fach", row.get('Fach'))
    add_metadata(document, "Institution", row.get('Institution'))

    document.add_paragraph("")

    # 6.3. Vollständige Transkription
    document.add_heading("Vollständige Transkription:", level=3)
    main_text = str(row.get('Text_Full', 'KEIN HAUPTTEXT FÜR DIESEN EINTRAG GEFUNDEN'))
    document.add_paragraph(main_text)

    # --- Ende des Briefinhalts ---

    # 6.4. Seitenumbruch einfügen (außer beim letzten Eintrag)
    if index < len(df) - 1:
        document.add_section(WD_SECTION.NEW_PAGE)

# 7. Dokument speichern
document.save(OUTPUT_FILE_PATH)  # <-- Wichtig: Speichert nun in den Input-Ordner!
print(f"\n--- ERFOLG ---")
print(
    f"Das Dokument '{OUTPUT_FILENAME_BASE}' wurde erfolgreich mit {len(df)} chronologisch sortierten Seiten in Ihrem Ordner '{INPUT_FOLDER}' gespeichert.")