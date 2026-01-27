import csv
import os
import re
from datetime import datetime
from docx import Document
from docx.oxml import OxmlElement
from copy import deepcopy


def erstelle_brief_anhang():
    print("--- Brief-Anhang Generator (Reduzierte Version) ---")
    input_csv = input("Bitte den Pfad zur CSV-Datei eingeben: ").strip().strip('"')

    template_path = r"D:\heiBOX\Seafile\Masterarbeit_Ablage\anhang_heibox\template_appendix_letter.docx"
    output_folder = r"D:\heiBOX\Seafile\Masterarbeit_Ablage\anhang_heibox"

    zeitstempel = datetime.now().strftime("%y%m%d_%H%M")
    output_filename = f"{zeitstempel}_Anhang_Briefe.docx"
    output_docx = os.path.join(output_folder, output_filename)

    if not os.path.exists(input_csv) or not os.path.exists(template_path):
        print("Fehler: Dateien nicht gefunden! Bitte Pfade prüfen.")
        return

    doc = Document(template_path)

    # Vorlage extrahieren
    template_para = doc.paragraphs[0]
    template_table = doc.tables[0]
    template_para_xml = deepcopy(template_para._element)
    template_table_xml = deepcopy(template_table._element)

    # Reduzierte Liste der Platzhalter
    spalten_liste = [
        "BRIEF-TITEL", "ABS-ORT", "EMP-ORT",
        "ARCHIV", "TRANSK", "TAG-FUNK", "TAG-THEMA"
    ]

    try:
        with open(input_csv, mode='r', encoding='utf-8-sig') as csvfile:
            sample = csvfile.read(4096)
            csvfile.seek(0)
            dialect = csv.Sniffer().sniff(sample)
            reader = csv.DictReader(csvfile, dialect=dialect)

            # Vorlage im Dokument löschen
            template_para._element.getparent().remove(template_para._element)
            template_table._element.getparent().remove(template_table._element)

            for i, row in enumerate(reader, start=1):
                # 1. TITEL-ABSATZ einfügen
                new_p_element = deepcopy(template_para_xml)
                doc._element.body.append(new_p_element)
                aktiver_absatz = doc.paragraphs[-1]

                if "BRIEF-TITEL" in aktiver_absatz.text:
                    wert = str(row.get("BRIEF-TITEL", "")).strip()
                    wert = wert.replace('\u00A0', ' ')
                    aktiver_absatz.text = f"{i}. {wert}"

                # 2. TABELLE einfügen
                new_tbl_element = deepcopy(template_table_xml)
                doc._element.body.append(new_tbl_element)
                aktuelle_tabelle = doc.tables[-1]

                for zeile in aktuelle_tabelle.rows:
                    for zelle in zeile.cells:
                        for spalte in spalten_liste:
                            if spalte in zelle.text:
                                wert = str(row.get(spalte, "")).strip()
                                wert = wert.replace('\u00A0', ' ')
                                zelle.text = zelle.text.replace(spalte, wert)

                # 3. LEERZEILE einfügen
                doc._element.body.append(OxmlElement('w:p'))

        # --- BEREINIGUNG: DOPPELTE LEERZEICHEN & ZEILENABSTAND ---
        def clean_spaces(text):
            return re.sub(r' {2,}', ' ', text)

        for para in doc.paragraphs:
            para.text = clean_spaces(para.text)
            para.paragraph_format.line_spacing = 1.15

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        para.text = clean_spaces(para.text)
                        para.paragraph_format.line_spacing = 1.15

        doc.save(output_docx)
        print(f"\nErfolg! {i} Briefe wurden verarbeitet.")
        print(f"Datei erstellt: {output_docx}")

    except Exception as e:
        print(f"Ein Fehler ist aufgetreten: {e}")


if __name__ == "__main__":
    erstelle_brief_anhang()